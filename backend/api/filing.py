"""
Filing API — Strategic Filing Drafter endpoints
"""
import io
import re
import uuid
from typing import Dict, Any, Optional, List
from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, Form
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy.ext.asyncio import AsyncSession

from backend.core.dependencies import get_db, get_current_user, CurrentUser
from backend.core.rbac import require_permission
from backend.services.filing_service import get_filing_service
from backend.services.docx_service import generate_filing_docx
from backend.core.logger import get_logger

logger = get_logger(__name__)

router = APIRouter(prefix="/api/v1/filing", tags=["filing"])
filing_service = get_filing_service()


class SelectedCitationItem(BaseModel):
    case_name: str
    citation: Optional[str] = None
    court: str
    year: int
    source_url: Optional[str] = None


class GenerateFilingRequest(BaseModel):
    filing_type: str
    objective: str
    petitioner: str
    respondent: str
    court: str
    sections: Optional[str] = None
    facts: str
    relief: Optional[str] = None
    selected_citations: Optional[List[SelectedCitationItem]] = None


@router.post("/generate", response_model=Dict[str, Any])
@require_permission("create_drafts")
async def generate_filing(
    body: GenerateFilingRequest,
    current_user: CurrentUser = Depends(get_current_user),
    session: AsyncSession = Depends(get_db),
) -> Dict[str, Any]:
    """Generate a strategic court filing."""
    try:
        valid_objectives = [
            "Win on merits",
            "Delay proceedings",
            "Challenge jurisdiction",
            "Seek settlement",
            "Preserve appeal rights",
        ]

        if body.objective not in valid_objectives:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Invalid objective. Must be one of: {', '.join(valid_objectives)}",
            )

        client_name = f"{body.petitioner} vs {body.respondent}"
        facts = body.facts
        if body.sections:
            facts = f"{body.facts}\n\nAdditional sections: {body.sections}"

        # Build user-selected citation strings to pass to the LLM
        user_selected_citations: List[str] = []
        if body.selected_citations:
            for c in body.selected_citations:
                if c.citation:
                    user_selected_citations.append(f"{c.case_name} — {c.citation}")
                else:
                    user_selected_citations.append(c.case_name)

        result = await filing_service.generate_filing(
            firm_id=str(current_user.firm_id),
            filing_type=body.filing_type,
            objective=body.objective,
            court=body.court,
            client_name=client_name,
            facts=facts,
            relief=body.relief or "",
            session=session,
            user_selected_citations=user_selected_citations if user_selected_citations else None,
        )

        if not result["success"]:
            return {
                "success": False,
                "error": result.get("error", "Draft generation failed"),
                "blocking_issues": result.get("blocking_issues", []),
                "draft": result.get("draft"),
                "quality_scores": result.get("quality_scores"),
                "citation_verification": result.get("citation_verification"),
            }

        return {
            "success": True,
            "draft_id": result["draft_id"],
            "draft": result["draft"],
            "quality_scores": result["quality_scores"],
            "citation_verification": result["citation_verification"],
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to generate filing: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to generate filing",
        )


# ---------------------------------------------------------------------------
# Fill-in-the-blanks TEMPLATE flow (new Draft-a-Filing)
# ---------------------------------------------------------------------------

class TemplateCitationItem(BaseModel):
    case_name: str
    citation: Optional[str] = None


class GenerateTemplateRequest(BaseModel):
    court: str = ""
    input_text: str
    selected_citations: Optional[List[TemplateCitationItem]] = None


class BriefCheckRequest(BaseModel):
    brief: str
    court: str = ""
    type_hint: str = ""


@router.post("/brief-check", response_model=Dict[str, Any])
@require_permission("create_drafts")
async def brief_check(
    body: BriefCheckRequest,
    current_user: CurrentUser = Depends(get_current_user),
) -> Dict[str, Any]:
    """Pre-flight completeness score + clarifying questions for a drafting brief.
    Cheap (gpt-4o-mini); called live as the lawyer types."""
    try:
        result = await filing_service.analyze_brief(
            firm_id=str(current_user.firm_id),
            brief=body.brief,
            court=body.court,
            type_hint=body.type_hint,
        )
        return {"success": True, **result}
    except ValueError as e:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=str(e))
    except Exception as e:
        logger.error(f"Failed to analyze brief: {str(e)}")
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="Failed to check brief")


@router.post("/template", response_model=Dict[str, Any])
@require_permission("create_drafts")
async def generate_template(
    body: GenerateTemplateRequest,
    current_user: CurrentUser = Depends(get_current_user),
    session: AsyncSession = Depends(get_db),
) -> Dict[str, Any]:
    """Generate a fill-in-the-blanks filing template from a free-text description."""
    try:
        result = await filing_service.generate_template(
            firm_id=str(current_user.firm_id),
            court=body.court,
            input_text=body.input_text,
            selected_citations=[c.model_dump() for c in (body.selected_citations or [])],
            is_existing_draft=False,
        )
        return {"success": True, **result}
    except ValueError as e:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=str(e))
    except Exception as e:
        logger.error(f"Failed to generate filing template: {str(e)}")
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="Failed to generate template")


@router.post("/template/upload", response_model=Dict[str, Any])
@require_permission("create_drafts")
async def generate_template_from_file(
    file: UploadFile = File(...),
    court: str = Form(""),
    current_user: CurrentUser = Depends(get_current_user),
    session: AsyncSession = Depends(get_db),
) -> Dict[str, Any]:
    """Upload a pre-written draft (PDF/DOCX); rewrite + templatize it."""
    ext = (file.filename or "").rsplit(".", 1)[-1].lower()
    if ext not in ("pdf", "docx"):
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Only PDF and DOCX files are supported.")
    file_bytes = await file.read()
    if len(file_bytes) > 50 * 1024 * 1024:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="File exceeds 50 MB limit.")
    try:
        text = await filing_service.extract_text(file_bytes, ext)
        if not text.strip():
            raise ValueError("No selectable text found — the file may be a scanned image.")
        result = await filing_service.generate_template(
            firm_id=str(current_user.firm_id),
            court=court,
            input_text=text,
            selected_citations=[],
            is_existing_draft=True,
        )
        return {"success": True, **result}
    except ValueError as e:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=str(e))
    except Exception as e:
        logger.error(f"Failed to templatize uploaded draft: {str(e)}")
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="Failed to process the uploaded draft")


class ExportTemplateRequest(BaseModel):
    title: str = "Draft"
    filled_markdown: str
    court: str = ""


def _markdown_to_docx(title: str, text: str) -> bytes:
    """Render the filled draft markdown to a .docx (headings + paragraphs)."""
    from docx import Document as DocxDocument
    from docx.shared import Pt

    doc = DocxDocument()
    if title:
        doc.add_heading(title, level=1)
    for raw in (text or "").split("\n"):
        line = raw.rstrip()
        if not line.strip():
            doc.add_paragraph("")
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        else:
            p = doc.add_paragraph(re.sub(r"\*\*([^*]+)\*\*", r"\1", line))
            if p.runs:
                p.runs[0].font.size = Pt(11)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@router.post("/template/export")
@require_permission("create_drafts")
async def export_template(
    body: ExportTemplateRequest,
    current_user: CurrentUser = Depends(get_current_user),
):
    """Export the filled-in template as a .docx."""
    try:
        docx_bytes = _markdown_to_docx(body.title, body.filled_markdown)
        return StreamingResponse(
            io.BytesIO(docx_bytes),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": 'attachment; filename="draft.docx"'},
        )
    except Exception as e:
        logger.error(f"Failed to export template: {str(e)}")
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="Failed to export draft")


@router.get("/{draft_id}", response_model=Dict[str, Any])
@require_permission("create_drafts")
async def get_draft(
    draft_id: str,
    current_user: CurrentUser = Depends(get_current_user),
    session: AsyncSession = Depends(get_db),
) -> Dict[str, Any]:
    """
    Get a draft by ID

    Args:
        draft_id: Draft ID
        current_user: Current authenticated user
        session: Database session

    Returns:
        Draft data with quality scores and citation badges
    """
    try:
        draft = await filing_service.get_draft(
            draft_id=draft_id,
            firm_id=str(current_user.firm_id),
            session=session,
        )

        if not draft:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Draft not found",
            )

        return {
            "success": True,
            "draft": draft,
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to get draft {draft_id}: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to retrieve draft",
        )


@router.post("/{draft_id}/regenerate", response_model=Dict[str, Any])
@require_permission("create_drafts")
async def regenerate_draft(
    draft_id: str,
    current_user: CurrentUser = Depends(get_current_user),
    session: AsyncSession = Depends(get_db),
) -> Dict[str, Any]:
    """
    Regenerate a draft with same inputs

    Args:
        draft_id: Existing draft ID
        current_user: Current authenticated user
        session: Database session

    Returns:
        New generated draft
    """
    try:
        result = await filing_service.regenerate_draft(
            draft_id=draft_id,
            firm_id=str(current_user.firm_id),
            session=session,
        )

        if not result["success"]:
            return {
                "success": False,
                "error": result.get("error", "Draft regeneration failed"),
                "blocking_issues": result.get("blocking_issues", []),
                "draft": result.get("draft"),
                "quality_scores": result.get("quality_scores"),
                "citation_verification": result.get("citation_verification"),
            }

        return {
            "success": True,
            "draft_id": result["draft_id"],
            "draft": result["draft"],
            "quality_scores": result["quality_scores"],
            "citation_verification": result["citation_verification"],
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to regenerate draft {draft_id}: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to regenerate draft",
        )


class ExportDraftRequest(BaseModel):
    """Body for DOCX export — sections passed from frontend state since DB stores combined text."""
    draft_sections: Dict[str, str]
    filing_type: str = "Filing"
    petitioner: str = ""
    respondent: str = ""
    court: str = ""
    citations_used: Optional[List[str]] = None


@router.post("/{draft_id}/export")
@require_permission("create_drafts")
async def export_draft(
    draft_id: str,
    body: ExportDraftRequest,
    current_user: CurrentUser = Depends(get_current_user),
    session: AsyncSession = Depends(get_db),
):
    """
    Export draft as a formatted .docx file.
    Sections are passed in the request body (they live in frontend state;
    DB stores only the combined plain-text content).
    """
    try:
        # Ownership check — 404 if draft not found or wrong firm
        draft = await filing_service.get_draft(
            draft_id=draft_id,
            firm_id=str(current_user.firm_id),
            session=session,
        )
        if not draft:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Draft not found",
            )

        docx_bytes = generate_filing_docx(
            draft_sections=body.draft_sections,
            filing_type=body.filing_type,
            petitioner=body.petitioner,
            respondent=body.respondent,
            court=body.court,
            citations_used=body.citations_used,
        )

        filename = f"draft_{draft_id[:8]}.docx"
        return StreamingResponse(
            io.BytesIO(docx_bytes),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to export draft {draft_id}: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to export draft",
        )