"""
Reply Service — Smart Reply Generator
Handles allegation extraction and reply generation from legal notices.

Flow:
  extract_allegations_from_bytes(file_bytes, file_ext, firm_id)
    1. Extract raw text from uploaded file (pypdf / python-docx)
    2. Sanitise text (GAP-001)
    3. Call GPT-4o-mini to extract structured allegations
    4. Return NoticeExtraction dict

  generate_reply(document_id, firm_id, allegation_responses, session)
    1. Load document OCR text for context
    2. Format lawyer's stance decisions
    3. Call GPT-4o-mini to generate full formal reply letter
    4. Store in law.drafts
    5. Return (draft_id, reply_text)

Rules:
- All AI calls through LLMService only
- Never log document content or PII (GAP-032)
- firm_id always from JWT — never trusted from caller outside service
"""
import io
import re
import json
import uuid
import asyncio
from typing import Optional, List, Dict, Any, Tuple
from datetime import datetime

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from backend.models.law_document import Document
from backend.models.law_draft import Draft, DraftType, DraftStatus
from backend.services.llm_service import get_llm_service, ModelType
from backend.services.prompts.reply_generator import allegation_extraction_prompt
from backend.core.logger import get_logger
from backend.core.sanitiser import sanitise_document_text

logger = get_logger(__name__)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _strip_fences(text: str) -> str:
    text = text.strip()
    if text.startswith("```"):
        text = re.sub(r"^```[a-z]*\n?", "", text)
        text = re.sub(r"\n?```$", "", text.strip())
    return text.strip()


def _parse_json(raw: str) -> dict:
    """Parse JSON from LLM response robustly.
    Handles: markdown fences, leading/trailing prose, truncated responses."""
    cleaned = _strip_fences(raw)

    # Direct parse
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        pass

    # Extract the first {...} block (handles prose before/after JSON)
    match = re.search(r'\{[\s\S]*\}', cleaned)
    if match:
        try:
            return json.loads(match.group())
        except json.JSONDecodeError:
            pass

    # Last resort: try to recover a truncated JSON by appending closing brackets
    # This handles cases where max_tokens cut the response mid-JSON
    for suffix in [']}', ']}]}', '}}', '}]}', ']}}}']:
        try:
            return json.loads(cleaned + suffix)
        except json.JSONDecodeError:
            pass

    logger.error(f"JSON parse failed in reply — raw={cleaned[:300]}")
    raise ValueError(f"Could not parse LLM response as JSON. Raw start: {cleaned[:100]}")


REPLY_SYSTEM_PROMPT = """You are an expert Indian lawyer drafting a formal reply to a legal notice for a client practising in Punjab/Haryana courts.
Generate a complete, professional reply notice in formal legal English.

Format the reply as a proper legal letter:
- Date line
- TO: [Sender's details from notice]
- SUB: Reply to Legal Notice dated [date if known]
- Opening paragraph: general denial of liability
- Numbered paragraphs, one per allegation — admit/deny/partially admit as instructed
- Closing paragraph: reserve all legal rights
- Advocate signature block placeholder

CRITICAL — how to phrase each paragraph:

ADMIT: State the fact cleanly. End the sentence there.
  Do NOT add commentary, inferences, qualifications, or context after the admitted fact.
  WRONG: "My client admits X, which is a matter of public record and contradicts Y."
  RIGHT: "My client admits that X."

DENY: State your client's POSITIVE CONTRARY POSITION. NEVER use double negatives.
  Always use "My client avers that..." — never "My client denies that..."

  MANDATORY REWRITE TABLE — apply these transformations when drafting DENY paragraphs:
  | Allegation contains...         | Draft must say...                                      |
  | "no complaint was made"        | "a complaint was duly made by the Pilot-in-Command"    |
  | "was not informed"             | "was duly informed of the complaints against him"      |
  | "no intervention by crew"      | "the cabin crew duly intervened"                       |
  | "was not constituted"          | "was duly constituted in accordance with the CAR"      |
  | "did not apologize"            | "did not in fact apologize" OR drop this framing       |
  | "nor was he uninformed"        | DELETE — replace with "was duly informed of..."        |
  | "denies that no"               | NEVER WRITE THIS — always rephrase as positive avers   |

  SELF-CHECK before writing each DENY paragraph: read your draft sentence.
  If it contains "not un-", "nor was he un-", "denies that no", or "denies that...not" — rewrite it.

PARTIAL: Identify the specific part admitted and the specific part denied, each stated clearly.

LEGAL PROVISIONS: If a legal_basis_claimed is provided for a denied allegation, address it directly.
  Always quote the specific paragraph/section number in your response (e.g. "Paragraph 3.0 of the CAR", "Section 138 NI Act").
  Either distinguish the provision, cite a counter-provision, or explain why it does not apply to the facts.

DOCUMENTARY SOURCES: If an allegation cites a specific government letter, official communication, or documentary source
  (e.g. "DGCA letter dated XX"), do NOT assert a contrary document that is not in the provided material.
  Instead, challenge the interpretation or legal consequence of that document.
  Example — if sender says "DGCA letter confirmed Committee not constituted", do NOT say "DGCA confirmed procedures were followed".
  Instead say: "My client avers that even assuming the contents of the said DGCA letter, the same does not vitiate the action taken, as..."

Never fabricate facts — only refer to events and actions explicitly present in the notice or the allegation decisions provided.
Never invent favorable documentary confirmations, witness reports, or official communications not present in the source material."""


class ReplyService:

    def __init__(self):
        self.llm = get_llm_service()

    # ------------------------------------------------------------------
    # Text extraction from file bytes
    # ------------------------------------------------------------------

    async def _extract_text(self, file_bytes: bytes, file_ext: str) -> str:
        ft = file_ext.lower().lstrip(".")
        try:
            if ft == "pdf":
                from pypdf import PdfReader
                reader = PdfReader(io.BytesIO(file_bytes))
                raw = "\n".join(page.extract_text() or "" for page in reader.pages).strip()
                if not raw:
                    raise ValueError(
                        "PDF appears to be a scanned image — no selectable text found. "
                        "Please upload a text-based PDF."
                    )
                return raw
            elif ft == "docx":
                from docx import Document as DocxDocument
                doc_obj = DocxDocument(io.BytesIO(file_bytes))
                return "\n".join(p.text for p in doc_obj.paragraphs).strip()
            else:
                raise ValueError(f"Unsupported file type '{ft}'. Supported: pdf, docx.")
        except ValueError:
            raise
        except Exception as exc:
            raise RuntimeError(f"Could not read file: {exc}") from exc

    # ------------------------------------------------------------------
    # Allegation extraction
    # ------------------------------------------------------------------

    async def _run_allegation_extraction(self, safe_text: str, firm_id: str) -> dict:
        # Truncate to ~12000 chars — legal notices can run 8-10 pages and the
        # strongest legal arguments (CAR violations, demands) are typically in
        # the latter paragraphs. 6000 chars only covered ~first 3 pages.
        truncated = safe_text[:12000]

        system_prompt = allegation_extraction_prompt.system_prompt
        user_prompt = allegation_extraction_prompt.format_user_prompt(notice_text=truncated)
        try:
            response_text = await asyncio.wait_for(
                self.llm.call_completion(
                    system_prompt=system_prompt,
                    user_prompt=user_prompt,
                    model=ModelType(allegation_extraction_prompt.model.value),
                    temperature=0.0,
                    max_tokens=3000,  # increased for longer notices with more allegations
                    firm_id=firm_id,
                    timeout=90,
                ),
                timeout=90.0,
            )
        except asyncio.TimeoutError:
            raise ValueError("AI request timed out — please try again")
        except Exception as exc:
            raise RuntimeError(f"Allegation extraction failed: {exc}") from exc

        data = _parse_json(response_text)
        # Ensure required keys with defaults
        data.setdefault("sender", None)
        data.setdefault("recipient", None)
        data.setdefault("notice_date", None)
        data.setdefault("notice_type", "other")
        data.setdefault("allegations", [])
        logger.info(f"Extracted {len(data['allegations'])} allegations firm={firm_id}")
        return data

    async def extract_allegations_from_bytes(
        self,
        file_bytes: bytes,
        file_ext: str,
        firm_id: str,
    ) -> dict:
        """Extract allegations from uploaded file — no DB document needed."""
        raw_text = await self._extract_text(file_bytes, file_ext)
        safe_text = sanitise_document_text(raw_text)
        return await self._run_allegation_extraction(safe_text, firm_id)

    async def extract_allegations(
        self,
        document_id: str,
        firm_id: str,
        session: AsyncSession,
    ) -> dict:
        """Extract allegations from an already-stored document."""
        stmt = select(Document).where(
            Document.id == uuid.UUID(document_id),
            Document.firm_id == uuid.UUID(firm_id),
            Document.deleted_at.is_(None),
        )
        result = await session.execute(stmt)
        document = result.scalar_one_or_none()
        if not document:
            raise ValueError(f"Document not found: {document_id}")
        if not document.ocr_text:
            raise ValueError("Document has no extractable text. Re-upload the file.")
        safe_text = sanitise_document_text(document.ocr_text)
        return await self._run_allegation_extraction(safe_text, firm_id)

    # ------------------------------------------------------------------
    # Reply generation
    # ------------------------------------------------------------------

    async def generate_reply(
        self,
        document_id: str,
        firm_id: str,
        allegation_responses: List[Dict[str, Any]],
        session: AsyncSession,
    ) -> Tuple[str, str]:
        """
        Generate complete formal reply.
        Returns (draft_id, reply_text).
        """
        # Load notice context from DB (optional — non-fatal if missing)
        notice_context = ""
        try:
            stmt = select(Document).where(
                Document.id == uuid.UUID(document_id),
                Document.firm_id == uuid.UUID(firm_id),
                Document.deleted_at.is_(None),
            )
            result = await session.execute(stmt)
            doc = result.scalar_one_or_none()
            if doc and doc.ocr_text:
                notice_context = sanitise_document_text(doc.ocr_text[:3000])
        except Exception:
            pass

        def _fmt_response(i: int, r: dict) -> str:
            parts = [
                f"Point {r.get('point_number', i + 1)}:",
                f"Stance={r.get('stance', 'deny').upper()}",
                f"Allegation: {r.get('allegation', '')}",
            ]
            if r.get('legal_basis_claimed'):
                parts.append(f"Sender's claimed legal basis: {r['legal_basis_claimed']}")
            if r.get('grounds'):
                parts.append(f"Lawyer's grounds: {r['grounds']}")
            return " | ".join(parts)

        responses_text = "\n".join([
            _fmt_response(i, r) for i, r in enumerate(allegation_responses)
        ])

        user_prompt = f"""Draft a complete formal reply to legal notice.

Notice content (first 3000 chars):
{notice_context or '[Notice text not available — draft based on allegation decisions below]'}

Lawyer's decisions on each allegation:
{responses_text}

Generate the complete reply notice text, properly formatted as a legal letter ready to send."""

        try:
            reply_text = await asyncio.wait_for(
                self.llm.call_completion(
                    system_prompt=REPLY_SYSTEM_PROMPT,
                    user_prompt=user_prompt,
                    model=ModelType.GPT4O_MINI,
                    temperature=0.0,
                    max_tokens=2500,
                    firm_id=firm_id,
                    timeout=90,
                ),
                timeout=90.0,
            )
        except asyncio.TimeoutError:
            raise ValueError("AI request timed out — please try again")
        except Exception as exc:
            raise RuntimeError(f"Reply generation failed: {exc}") from exc

        # Store draft
        draft = Draft(
            firm_id=uuid.UUID(firm_id),
            draft_type=DraftType.SMART_REPLY,
            title="Reply to Legal Notice",
            content=reply_text,
            status=DraftStatus.GENERATED,
        )
        session.add(draft)
        await session.flush()
        await session.commit()

        logger.info(f"Reply draft generated draft={draft.id} firm={firm_id}")
        return str(draft.id), reply_text

    # ------------------------------------------------------------------
    # DOCX export
    # ------------------------------------------------------------------

    async def export_reply_docx(
        self,
        draft_id: str,
        firm_id: str,
        session: AsyncSession,
    ) -> bytes:
        stmt = select(Draft).where(
            Draft.id == uuid.UUID(draft_id),
            Draft.firm_id == uuid.UUID(firm_id),
            Draft.deleted_at.is_(None),
        )
        result = await session.execute(stmt)
        draft = result.scalar_one_or_none()
        if not draft:
            raise ValueError(f"Draft not found: {draft_id}")

        from docx import Document as DocxDocument
        from docx.shared import Pt, RGBColor

        doc = DocxDocument()
        doc.add_heading("REPLY TO LEGAL NOTICE", level=1)
        doc.add_paragraph("")

        for line in draft.content.split("\n"):
            stripped = line.strip()
            if stripped:
                p = doc.add_paragraph(stripped)
                if p.runs:
                    p.runs[0].font.size = Pt(11)
            else:
                doc.add_paragraph("")

        doc.add_paragraph("")
        footer = doc.add_paragraph("Generated by Nikhar Legal Workspace")
        if footer.runs:
            footer.runs[0].font.size = Pt(8)
            footer.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()


# ---------------------------------------------------------------------------
# Singleton
# ---------------------------------------------------------------------------

_reply_service: Optional[ReplyService] = None


def get_reply_service() -> ReplyService:
    global _reply_service
    if _reply_service is None:
        _reply_service = ReplyService()
    return _reply_service
