"""
Document service — all business logic for document processing
Handles: upload, storage, OCR, chunking, embedding, indexing
"""
import os
import uuid
from typing import Optional, List, Tuple, Dict, Any
from datetime import datetime
import asyncio
import math
from io import BytesIO

from sqlalchemy import select, update
from sqlalchemy.ext.asyncio import AsyncSession

from backend.models import Document, DocumentStatus
from backend.core.logger import get_logger
from backend.core.config import settings
from backend.services.rag_service import get_rag_service

logger = get_logger(__name__)

# Token estimation (rough approximation: ~4 chars = 1 token)
TOKEN_ESTIMATE_RATIO = 4
CHUNK_SIZE_TOKENS = 500
OVERLAP_TOKENS = 50


class DocumentService:
    """Document processing service"""
    
    # File validation
    ALLOWED_FILE_TYPES = {'pdf', 'docx', 'doc', 'jpg', 'jpeg', 'png', 'tiff', 'tif'}
    MAX_FILE_SIZE_BYTES = 50 * 1024 * 1024  # 50MB
    
    def __init__(self):
        """Initialize service with Azure clients"""
        self._blob_client = None
        self._doc_intelligence_client = None
        self._search_client = None
        self._embeddings_client = None
        self.rag_service = get_rag_service()
    
    async def validate_upload(self, filename: str, file_size: int) -> Tuple[bool, Optional[str]]:
        """
        Validate file upload
        
        Args:
            filename: Name of file being uploaded
            file_size: Size in bytes
            
        Returns:
            Tuple of (is_valid, error_message)
        """
        # Check file type
        file_ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else None
        if not file_ext or file_ext not in self.ALLOWED_FILE_TYPES:
            return False, f"File type '.{file_ext}' not allowed. Supported: {', '.join(self.ALLOWED_FILE_TYPES)}"
        
        # Check file size
        if file_size > self.MAX_FILE_SIZE_BYTES:
            size_mb = file_size / (1024 * 1024)
            return False, f"File size {size_mb:.1f}MB exceeds maximum 50MB"
        
        return True, None
    
    async def save_to_blob_storage(
        self,
        file_content: bytes,
        filename: str,
        firm_id: str,
        file_type: str
    ) -> str:
        """
        Save file to Azure Blob Storage
        
        Args:
            file_content: File bytes
            filename: Original filename
            firm_id: Firm ID for organization
            file_type: File extension
            
        Returns:
            Blob path (for later retrieval)
        """
        try:
            # TODO: Implement actual Azure Blob Storage upload
            # For now, return a mock path
            blob_path = f"documents/{firm_id}/{uuid.uuid4()}/{filename}"
            logger.info(f"Saved to blob: {blob_path}")
            return blob_path
        except Exception as e:
            logger.error(f"Failed to save to blob storage: {str(e)}")
            raise
    
    async def store_document_in_db(
        self,
        session: AsyncSession,
        firm_id: str,
        matter_id: Optional[str],
        filename: str,
        file_type: str,
        file_size_bytes: int,
        blob_path: str,
        user_id: Optional[str]
    ) -> Document:
        """
        Create document record in database
        
        Args:
            session: Database session
            firm_id: Firm ID
            matter_id: Optional matter ID
            filename: File name
            file_type: File type/extension
            file_size_bytes: File size
            blob_path: Path in blob storage
            user_id: User who uploaded
            
        Returns:
            Created Document object
        """
        try:
            # Convert matter_id string to UUID if provided
            matter_uuid = None
            if matter_id:
                try:
                    import uuid as uuid_module
                    matter_uuid = uuid_module.UUID(matter_id)
                except (ValueError, TypeError):
                    logger.warning(f"Invalid matter_id: {matter_id}")
            
            document = Document(
                firm_id=uuid.UUID(firm_id),
                matter_id=matter_uuid,
                file_name=filename,
                file_type=file_type.lower(),
                file_size_bytes=file_size_bytes,
                blob_path=blob_path,
                status=DocumentStatus.PENDING,
                uploaded_by_user_id=uuid.UUID(user_id) if user_id else None,
                upload_source="web"
            )
            
            session.add(document)
            await session.flush()  # Get the ID before commit
            await session.commit()
            
            logger.info(f"Created document record: {document.id}")
            return document
            
        except Exception as e:
            await session.rollback()
            logger.error(f"Failed to store document in DB: {str(e)}")
            raise
    
    async def get_document(self, session: AsyncSession, document_id: str, firm_id: str) -> Optional[Document]:
        """Get document by ID, scoped to firm"""
        try:
            result = await session.execute(
                select(Document).where(
                    Document.id == uuid.UUID(document_id),
                    Document.firm_id == uuid.UUID(firm_id),
                    Document.deleted_at.is_(None),
                )
            )
            return result.scalar_one_or_none()
        except Exception as e:
            logger.error(f"Failed to get document: {str(e)}")
            return None
    
    async def update_document_status(
        self,
        session: AsyncSession,
        document_id: str,
        status: DocumentStatus,
        error_reason: Optional[str] = None
    ) -> bool:
        """Update document status"""
        try:
            query = (
                update(Document)
                .where(Document.id == uuid.UUID(document_id))
                .values(status=status, error_reason=error_reason)
            )
            await session.execute(query)
            await session.commit()
            logger.info(f"Updated document {document_id} status to {status}")
            return True
        except Exception as e:
            await session.rollback()
            logger.error(f"Failed to update document status: {str(e)}")
            return False
    
    async def extract_text_from_document(
        self,
        file_content: bytes,
        file_type: str
    ) -> str:
        """
        Extract text from document.
        PDF/DOCX: extracted locally with pypdf / python-docx.
        Images and scanned PDFs: requires Azure Document Intelligence (TODO).
        """
        import io
        ft = file_type.lower().lstrip(".")

        try:
            if ft == "pdf":
                from pypdf import PdfReader
                reader = PdfReader(io.BytesIO(file_content))
                pages = [page.extract_text() or "" for page in reader.pages]
                text = "\n".join(pages).strip()
                if not text:
                    raise ValueError(
                        "PDF appears to be a scanned image — no selectable text found. "
                        "Azure Document Intelligence OCR is required for scanned documents."
                    )
                logger.info(f"Extracted {len(text)} chars from PDF ({len(file_content)} bytes)")
                return text

            elif ft in ("docx",):
                from docx import Document as DocxDocument
                doc = DocxDocument(io.BytesIO(file_content))
                text = "\n".join(p.text for p in doc.paragraphs).strip()
                logger.info(f"Extracted {len(text)} chars from DOCX ({len(file_content)} bytes)")
                return text

            else:
                raise ValueError(
                    f"Unsupported file type '{ft}' for text extraction. "
                    "Supported: pdf, docx."
                )

        except ValueError:
            raise
        except Exception as e:
            logger.error(f"Failed to extract text from {ft}: {e}")
            raise
    
    def chunk_text(self, text: str, chunk_size_tokens: int = CHUNK_SIZE_TOKENS, 
                   overlap_tokens: int = OVERLAP_TOKENS) -> List[str]:
        """
        Chunk text into overlapping chunks
        
        Args:
            text: Full text to chunk
            chunk_size_tokens: Target chunk size in tokens
            overlap_tokens: Overlap between chunks in tokens
            
        Returns:
            List of text chunks
        """
        try:
            # Convert tokens to approximate character count
            chunk_size_chars = int(chunk_size_tokens * TOKEN_ESTIMATE_RATIO)
            overlap_chars = int(overlap_tokens * TOKEN_ESTIMATE_RATIO)
            
            chunks = []
            start = 0
            
            while start < len(text):
                end = min(start + chunk_size_chars, len(text))
                chunk = text[start:end].strip()
                
                if chunk:  # Only add non-empty chunks
                    chunks.append(chunk)
                
                # Move to next position with overlap
                start = end - overlap_chars
                if start < len(text) and start + chunk_size_chars >= len(text):
                    # Ensure we get the last chunk
                    start = len(text) - chunk_size_chars
            
            logger.info(f"Created {len(chunks)} chunks from text ({len(text)} chars)")
            return chunks
            
        except Exception as e:
            logger.error(f"Failed to chunk text: {str(e)}")
            raise
    
    async def ingest_document(
        self,
        text: str,
        document_id: str,
        firm_id: str,
        filename: str,
        vertical: str = "law",
    ) -> List[Dict[str, Any]]:
        """Chunk, embed, and index a document text through the RAG pipeline."""
        try:
            chunks = self.rag_service.chunk_document(text, document_id, firm_id)
            chunks = await self.rag_service.embed_chunks(chunks)
            await self.rag_service.index_chunks(
                chunks=chunks,
                document_id=document_id,
                firm_id=firm_id,
                vertical=vertical,
                document_name=filename,
            )
            return chunks
        except Exception as e:
            logger.error(f"Failed to ingest document through RAG pipeline: {str(e)}")
            raise
    
    async def generate_embeddings(self, texts: List[str]) -> List[List[float]]:
        """
        Generate embeddings using text-embedding-ada-002
        
        Args:
            texts: List of text chunks to embed
            
        Returns:
            List of embedding vectors (1536 dimensions each)
        """
        try:
            # TODO: Implement actual Azure OpenAI embeddings call
            # For now, return mock embeddings
            embeddings = []
            for text in texts:
                # Mock 1536-dimensional embedding
                embedding = [0.1 * (i % 10) for i in range(1536)]
                embeddings.append(embedding)
            
            logger.info(f"Generated {len(embeddings)} embeddings")
            return embeddings
            
        except Exception as e:
            logger.error(f"Failed to generate embeddings: {str(e)}")
            raise
    
    async def index_in_search(
        self,
        document_id: str,
        chunks: List[str],
        embeddings: List[List[float]],
        firm_id: str,
        filename: str
    ) -> bool:
        """
        Index document chunks in Azure AI Search
        
        Args:
            document_id: Document ID
            chunks: Text chunks
            embeddings: Embedding vectors
            firm_id: Firm ID for filtering
            filename: Original filename
            
        Returns:
            Success status
        """
        try:
            # TODO: Implement actual Azure AI Search indexing
            # For now, just log
            logger.info(
                f"Indexed document {document_id} with {len(chunks)} chunks in "
                f"search index for firm {firm_id}"
            )
            return True
            
        except Exception as e:
            logger.error(f"Failed to index in search: {str(e)}")
            raise
    
    async def update_document_indexed(
        self,
        session: AsyncSession,
        document_id: str,
        chunk_count: int,
        ocr_text: Optional[str] = None,
    ) -> bool:
        """Update document as indexed"""
        try:
            query = (
                update(Document)
                .where(Document.id == uuid.UUID(document_id))
                .values(
                    status=DocumentStatus.INDEXED,
                    chunk_count=chunk_count,
                    indexed_at=datetime.utcnow(),
                    ocr_text=ocr_text,
                )
            )
            await session.execute(query)
            await session.commit()
            logger.info(f"Document {document_id} indexed with {chunk_count} chunks")
            return True
        except Exception as e:
            await session.rollback()
            logger.error(f"Failed to update document indexed status: {str(e)}")
            return False

    async def delete_document(
        self,
        session: AsyncSession,
        document_id: str,
        firm_id: str,
    ) -> bool:
        """
        Soft delete a document and remove from RAG index

        Args:
            session: Database session
            document_id: Document ID to delete
            firm_id: Firm ID for ownership verification

        Returns:
            Success status

        Raises:
            ValueError: If document not found or doesn't belong to firm
        """
        try:
            # Verify document exists and belongs to firm
            document = await self.get_document(session, document_id)
            if not document:
                raise ValueError(f"Document not found: {document_id}")

            if str(document.firm_id) != firm_id:
                raise ValueError(f"Document {document_id} does not belong to firm {firm_id}")

            # Soft delete from database
            query = (
                update(Document)
                .where(Document.id == uuid.UUID(document_id))
                .values(deleted_at=datetime.utcnow())
            )
            await session.execute(query)

            # Remove from RAG index
            await self.rag_service.delete_document_chunks(document_id, firm_id)

            await session.commit()

            # Log audit event
            logger.info(f"Document {document_id} soft deleted by firm {firm_id}")

            return True

        except ValueError:
            # Re-raise validation errors
            raise
        except Exception as e:
            await session.rollback()
            logger.error(f"Failed to delete document {document_id}: {str(e)}")
            raise


# Singleton instance
_document_service = None


def get_document_service() -> DocumentService:
    """Get or create document service instance"""
    global _document_service
    if _document_service is None:
        _document_service = DocumentService()
    return _document_service
