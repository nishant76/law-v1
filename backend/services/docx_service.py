"""
DOCX Service — generates formatted Word documents from draft sections
Uses python-docx to produce court-filing-ready .docx files
"""
import io
from typing import Dict, Optional
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from backend.core.logger import get_logger

logger = get_logger(__name__)


def _set_cell_border(cell, **kwargs):
    """Helper — not needed here but kept for future table use."""
    pass


def _add_horizontal_rule(doc: Document):
    """Add a thin horizontal line between sections."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def generate_filing_docx(
    draft_sections: Dict[str, str],
    filing_type: str,
    petitioner: str,
    respondent: str,
    court: str,
    citations_used: Optional[list] = None,
) -> bytes:
    """
    Generate a formatted DOCX from draft sections.

    Args:
        draft_sections: Dict with keys court_heading, parties_section,
                        facts_section, grounds_section, prayer_section, verification
        filing_type: e.g. "Bail Application"
        petitioner: Petitioner name
        respondent: Respondent name
        court: Court name
        citations_used: List of citation strings

    Returns:
        Raw bytes of the .docx file
    """
    doc = Document()

    # ── Page margins — A4, court-standard 1 inch on all sides ─────────────────
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # ── Default paragraph style ────────────────────────────────────────────────
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    def _para(
        text: str,
        bold: bool = False,
        italic: bool = False,
        underline: bool = False,
        align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
        size: int = 12,
        space_before: int = 0,
        space_after: int = 6,
        all_caps: bool = False,
    ):
        p = doc.add_paragraph()
        p.paragraph_format.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.underline = underline
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        if all_caps:
            run.font.all_caps = True
        return p

    # ── Section label helper ───────────────────────────────────────────────────
    def _section_heading(label: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(label)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        run.font.all_caps = True
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # ── 1. Court heading ───────────────────────────────────────────────────────
    court_heading = draft_sections.get("court_heading", "").strip()
    if court_heading:
        for line in court_heading.split("\n"):
            line = line.strip()
            if not line:
                continue
            _para(
                line,
                bold=True,
                underline=True,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                size=13,
                space_before=0,
                space_after=4,
            )
    else:
        _para(
            court.upper(),
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            size=13,
        )

    doc.add_paragraph()  # spacer

    # ── 2. Parties section ─────────────────────────────────────────────────────
    parties = draft_sections.get("parties_section", "").strip()
    if parties:
        for line in parties.split("\n"):
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue
            p = doc.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(line)
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
    else:
        _para(
            f"{petitioner}\n...Petitioner",
            align=WD_ALIGN_PARAGRAPH.CENTER,
            size=11,
        )
        _para("Versus", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=11)
        _para(
            f"{respondent}\n...Respondent",
            align=WD_ALIGN_PARAGRAPH.CENTER,
            size=11,
        )

    _add_horizontal_rule(doc)

    # ── 3. Facts ───────────────────────────────────────────────────────────────
    facts = draft_sections.get("facts_section", "").strip()
    if facts:
        _section_heading("Facts")
        p = doc.add_paragraph(facts)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(8)
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

    # ── 4. Grounds ────────────────────────────────────────────────────────────
    grounds = draft_sections.get("grounds_section", "").strip()
    if grounds:
        _section_heading("Grounds")
        p = doc.add_paragraph(grounds)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(8)
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

    # ── 5. Prayer ─────────────────────────────────────────────────────────────
    prayer = draft_sections.get("prayer_section", "").strip()
    if prayer:
        _section_heading("Prayer")
        p = doc.add_paragraph(prayer)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(8)
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

    # ── 6. Citations used ─────────────────────────────────────────────────────
    if citations_used:
        _add_horizontal_rule(doc)
        _section_heading("Citations Relied Upon")
        for c in citations_used:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(c)
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)

    # ── 7. Verification ───────────────────────────────────────────────────────
    verification = draft_sections.get("verification", "").strip()
    if verification:
        _add_horizontal_rule(doc)
        _section_heading("Verification")
        p = doc.add_paragraph(verification)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            run.italic = True

    # ── Signature block ───────────────────────────────────────────────────────
    doc.add_paragraph()
    _para(
        "Place: _______________",
        size=11,
        space_before=12,
        space_after=4,
    )
    _para("Date:  _______________", size=11, space_after=20)
    _para(
        "Counsel for the Petitioner",
        bold=True,
        align=WD_ALIGN_PARAGRAPH.RIGHT,
        size=11,
        space_before=24,
    )

    # Serialise to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    logger.info(f"Generated DOCX for {filing_type}: {len(draft_sections)} sections")
    return buf.read()
