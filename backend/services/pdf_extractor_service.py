"""
PDF Extractor Service — universal single-call extraction pipeline.

Extracts structured intelligence from any document type in a single LLM call.
Works for both indexed documents (from library) and direct uploads.

Rules:
- All AI calls through LLMService only
- firm_id from JWT — never trusted from caller outside service
- Never log document content or PII (GAP-032)
- Amber threshold: < 75% confidence (CLAUDE.md Feature #5)
"""
import io
import json
import re
import uuid
import asyncio
from typing import Optional, AsyncGenerator

from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select

from backend.models.law_document import Document, DocumentStatus
from backend.services.llm_service import get_llm_service, LLMService, ModelType as LLMModelType
from backend.services.prompts.pdf_extractor import (
    pdf_extractor_prompt,
    READABLE_SYSTEM_PROMPT,
    READABLE_USER_TEMPLATE,
    READABLE_MODEL,
)
from backend.services.pdf_extractor_validator import validate_and_correct
from backend.core.logger import get_logger
from backend.core.sanitiser import sanitise_document_text

logger = get_logger(__name__)


def _strip_fences(text: str) -> str:
    text = text.strip()
    if text.startswith("```"):
        text = re.sub(r"^```[a-z]*\n?", "", text)
        text = re.sub(r"\n?```$", "", text.strip())
    return text.strip()


def _parse_llm_json(raw: str, context: str) -> dict:
    cleaned = _strip_fences(raw)
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError as exc:
        logger.error(f"JSON parse error in {context}: {exc} — raw={cleaned[:200]}")
        raise ValueError(f"LLM returned non-JSON response ({context}): {exc}") from exc


class PDFExtractorService:

    def __init__(self, llm_service: LLMService):
        self._llm = llm_service

    async def _call_extraction(self, text: str, firm_id: str) -> dict:
        user_prompt = pdf_extractor_prompt.format_user_prompt(document_text=text)
        try:
            raw = await asyncio.wait_for(
                self._llm.call_completion(
                    system_prompt=pdf_extractor_prompt.system_prompt,
                    user_prompt=user_prompt,
                    model=LLMModelType(pdf_extractor_prompt.model.value),
                    temperature=0.0,
                    max_tokens=pdf_extractor_prompt.max_tokens,
                    firm_id=firm_id,
                    timeout=90,
                ),
                timeout=90.0,
            )
        except asyncio.TimeoutError:
            raise RuntimeError("AI request timed out — please try again")
        except Exception as exc:
            logger.error(f"LLM call failed: {exc}")
            raise RuntimeError(f"Extraction failed: {exc}") from exc
        parsed = _parse_llm_json(raw, "universal_extract")
        return validate_and_correct(parsed, document_text=text)

    async def extract_fields(
        self,
        db: AsyncSession,
        document_id: str,
        firm_id: str,
    ) -> dict:
        stmt = select(Document).where(
            Document.id == uuid.UUID(document_id),
            Document.deleted_at.is_(None),
        )
        result = await db.execute(stmt)
        doc: Optional[Document] = result.scalar_one_or_none()

        if doc is None or str(doc.firm_id) != firm_id:
            raise ValueError("Document not found")

        if doc.status != DocumentStatus.INDEXED:
            raise ValueError(
                f"Document is not yet indexed (status={doc.status.value}). "
                "Please wait for indexing to complete before extracting fields."
            )

        if not doc.ocr_text:
            raise ValueError("Document has no extracted text. Re-upload or check if OCR failed.")

        safe_text = sanitise_document_text(doc.ocr_text)
        parsed = await self._call_extraction(safe_text, firm_id)
        return {"document_id": document_id, **parsed}

    async def extract_fields_from_bytes(
        self,
        file_bytes: bytes,
        file_ext: str,
        firm_id: str,
    ) -> dict:
        ft = file_ext.lower().lstrip(".")
        try:
            if ft == "pdf":
                from pypdf import PdfReader
                reader = PdfReader(io.BytesIO(file_bytes))
                pages = [page.extract_text() or "" for page in reader.pages]
                raw_text = "\n".join(pages).strip()
                if not raw_text:
                    raise ValueError(
                        "PDF appears to be a scanned image — no selectable text found."
                    )
            elif ft == "docx":
                from docx import Document as DocxDocument
                doc_obj = DocxDocument(io.BytesIO(file_bytes))
                raw_text = "\n".join(p.text for p in doc_obj.paragraphs).strip()
            else:
                raise ValueError(f"Unsupported file type '{ft}'. Supported: pdf, docx.")
        except ValueError:
            raise
        except Exception as exc:
            logger.error(f"Text extraction failed for {ft}: {exc}")
            raise RuntimeError(f"Could not read file: {exc}") from exc

        logger.info(f"Direct extraction: {ft}, {len(file_bytes)} bytes, firm={firm_id}")
        safe_text = sanitise_document_text(raw_text)
        try:
            parsed = await self._call_extraction(safe_text, firm_id)
        except asyncio.TimeoutError:
            raise RuntimeError("AI request timed out — please try again")
        except ValueError as exc:
            raise RuntimeError(str(exc)) from exc
        except Exception as exc:
            logger.error(f"Extraction failed: {exc}")
            raise RuntimeError(f"Field extraction failed: {exc}") from exc

        return {"document_id": "direct", **parsed}, safe_text


    async def extract_fields_stream(
        self,
        file_bytes: bytes,
        file_ext: str,
        firm_id: str,
    ) -> AsyncGenerator[dict, None]:
        """
        Two-phase streaming extraction:

        Phase 1 — readable markdown prose streams token by token so the
                   lawyer can start reading immediately.
        Phase 2 — full JSON extraction runs after Phase 1 completes;
                   the 'result' event carries the structured data.

        SSE event types:
          progress  {"stage": "reading"|"analysing"|"extracting", "message": str}
          token     {"text": str}   ← Phase 1 markdown tokens
          result    {"data": {...}, "raw_text": str}
          (errors raised as exceptions — caller emits error SSE)
        """
        ft = file_ext.lower().lstrip(".")

        # ── 1. Extract raw text from file ─────────────────────────────────────
        yield {"type": "progress", "stage": "reading", "message": "Reading document…"}
        try:
            if ft == "pdf":
                from pypdf import PdfReader
                reader = PdfReader(io.BytesIO(file_bytes))
                pages = [page.extract_text() or "" for page in reader.pages]
                raw_text = "\n".join(pages).strip()
                if not raw_text:
                    raise ValueError(
                        "PDF appears to be a scanned image — no selectable text found."
                    )
            elif ft == "docx":
                from docx import Document as DocxDocument
                doc_obj = DocxDocument(io.BytesIO(file_bytes))
                raw_text = "\n".join(p.text for p in doc_obj.paragraphs).strip()
            else:
                raise ValueError(f"Unsupported file type '{ft}'.")
        except ValueError:
            raise
        except Exception as exc:
            logger.error(f"Text extraction failed for {ft}: {exc}")
            raise RuntimeError(f"Could not read file: {exc}") from exc

        safe_text = sanitise_document_text(raw_text)

        # ── 2. Phase 1: stream a readable markdown analysis ───────────────────
        yield {"type": "progress", "stage": "analysing", "message": "Analysing document…"}

        readable_user_prompt = READABLE_USER_TEMPLATE.format(
            document_text=safe_text[:80_000]  # ~60 pages — covers any district court document
        )

        async for chunk in self._llm.call_completion_stream(
            system_prompt=READABLE_SYSTEM_PROMPT,
            user_prompt=readable_user_prompt,
            model=LLMModelType(READABLE_MODEL.value),
            temperature=0.0,
            max_tokens=6000,  # complex judgments with witnesses/exhibits need more room
            firm_id=firm_id,
        ):
            yield {"type": "token", "text": chunk}

        # Phase 1 complete — the streamed markdown IS the result.
        # No second LLM call, no JSON extraction, no restructuring.
        yield {"type": "result", "raw_text": raw_text}


_pdf_extractor_service: Optional[PDFExtractorService] = None


def get_pdf_extractor_service() -> PDFExtractorService:
    global _pdf_extractor_service
    if _pdf_extractor_service is None:
        _pdf_extractor_service = PDFExtractorService(llm_service=get_llm_service())
    return _pdf_extractor_service
