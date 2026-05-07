"""
Case Synopsis Service
Business logic for Case Synopsis Generator (Feature #6 in CLAUDE.md).

Flow:
  generate_synopsis(document_id, firm_id)
    1. Load document — verify firm ownership (return 404 if wrong firm)
    2. Sanitise OCR text before injecting into prompt (GAP-001)
    3. Call GPT-4o-mini via LLMService
    4. Parse JSON response
    5. Store result in law.drafts (draft_type=CASE_SYNOPSIS)
    6. Return SynopsisResult

  export_synopsis_docx(synopsis_id, firm_id)
    1. Load draft — verify firm ownership
    2. Build formatted .docx from structured content
    3. Return bytes for HTTP response

Rules:
- All AI calls through LLMService only
- firm_id always from JWT — never trusted from caller outside this service
- Soft deletes only — never hard delete
- Never log document content or PII (GAP-032)
"""
import json
import re
import uuid
import asyncio
from io import BytesIO
from typing import Optional
from datetime import datetime

from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select

from backend.models.law_document import Document, DocumentStatus
from backend.models.law_draft import Draft, DraftType, DraftStatus
from backend.services.llm_service import get_llm_service, LLMService, ModelType
from backend.services.prompts.case_synopsis import case_synopsis_prompt
from backend.core.logger import get_logger
from backend.core.sanitiser import sanitise_document_text
from backend.core.sanitiser import sanitise_document_text

logger = get_logger(__name__)


# ---------------------------------------------------------------------------
# Result dataclass (no ORM dependency in return types)
# ---------------------------------------------------------------------------

class SynopsisResult:
    """Structured output from synopsis generation"""

    def __init__(
        self,
        synopsis_id: str,
        document_id: str,
        case_name: Optional[str],
        petitioner: Optional[str],
        respondent: Optional[str],
        court: Optional[str],
        judgment_date: Optional[str],
        case_number: Optional[str],
        facts: Optional[str],
        issues: list,
        held: Optional[str],
        citations_used: list,
        relief_granted: Optional[str],
        confidence: int,
        created_at: datetime,
    ):
        self.synopsis_id = synopsis_id
        self.document_id = document_id
        self.case_name = case_name
        self.petitioner = petitioner
        self.respondent = respondent
        self.court = court
        self.judgment_date = judgment_date
        self.case_number = case_number
        self.facts = facts
        self.issues = issues
        self.held = held
        self.citations_used = citations_used
        self.relief_granted = relief_granted
        self.confidence = confidence
        self.created_at = created_at


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _parse_synopsis_json(raw_response: str) -> dict:
    """
    Parse LLM JSON response.
    Strips markdown code fences if model wraps output.
    Returns dict with all OUTPUT_SCHEMA keys, defaulting missing keys to None/[].
    """
    # Strip ```json ... ``` fences if present
    text = raw_response.strip()
    if text.startswith("```"):
        text = re.sub(r"^```[a-z]*\n?", "", text)
        text = re.sub(r"\n?```$", "", text.strip())
    text = text.strip()

    # Direct parse
    try:
        data = json.loads(text)
    except json.JSONDecodeError:
        # Try to extract first {...} block (handles prose around JSON)
        match = re.search(r'\{[\s\S]*\}', text)
        if match:
            try:
                data = json.loads(match.group())
            except json.JSONDecodeError:
                data = None
        else:
            data = None

        if data is None:
            logger.error(f"Failed to parse synopsis JSON — raw={text[:300]}")
            raise ValueError(f"LLM returned non-JSON response: {text[:100]}")

    # Normalise — ensure all expected keys exist
    defaults = {
        "document_type": "other",
        "case_name": None,
        "petitioner": None,
        "respondent": None,
        "court": None,
        "judgment_date": None,
        "case_number": None,
        "facts": None,
        "issues": [],
        "held": None,
        "citations_used": [],
        "relief_granted": None,
        "confidence": 0,
    }
    defaults.update(data)
    return defaults


def _build_synopsis_title(parsed: dict) -> str:
    """Derive a human-readable title for the draft record"""
    case_name = parsed.get("case_name") or "Untitled Case"
    return f"Case Synopsis — {case_name}"


# Document types accepted by Case Synopsis Generator
_SYNOPSIS_ALLOWED_TYPES = {"judgment", "petition", "order"}


def _validate_document_type(parsed: dict) -> None:
    """
    Raise ValueError with a structured prefix if document is not suitable
    for synopsis generation (i.e. not a court judgment, petition, or order).

    Callers catch "WRONG_DOCUMENT_TYPE:<type>" and return a structured
    error response to the frontend.
    """
    doc_type = str(parsed.get("document_type") or "other").lower().strip().replace(" ", "_")
    if doc_type not in _SYNOPSIS_ALLOWED_TYPES:
        raise ValueError(f"WRONG_DOCUMENT_TYPE:{doc_type}")


# ---------------------------------------------------------------------------
# Main service
# ---------------------------------------------------------------------------

class SynopsisService:

    def __init__(self, llm_service: LLMService):
        self._llm = llm_service

    async def generate_synopsis(
        self,
        db: AsyncSession,
        document_id: str,
        firm_id: str,
        user_id: Optional[str] = None,
    ) -> SynopsisResult:
        """
        Generate a structured case synopsis from an already-indexed document.

        Args:
            db:          Async database session
            document_id: UUID of the document (law.documents table)
            firm_id:     Firm ID from JWT — used for ownership check
            user_id:     User ID from JWT — stored on draft for audit

        Returns:
            SynopsisResult with all extracted fields

        Raises:
            ValueError:  Document not found / not owned by firm / not indexed
            RuntimeError: LLM call failed
        """
        # ------------------------------------------------------------------
        # 1. Load document and verify firm ownership
        # ------------------------------------------------------------------
        stmt = select(Document).where(
            Document.id == uuid.UUID(document_id),
            Document.deleted_at.is_(None),
        )
        result = await db.execute(stmt)
        doc: Optional[Document] = result.scalar_one_or_none()

        if doc is None or str(doc.firm_id) != firm_id:
            # Return 404 not 403 — never reveal existence to wrong firm (CLAUDE.md security)
            raise ValueError("Document not found")

        if doc.status != DocumentStatus.INDEXED:
            raise ValueError(
                f"Document is not yet indexed (status={doc.status.value}). "
                "Please wait for indexing to complete before generating a synopsis."
            )

        if not doc.ocr_text:
            raise ValueError(
                "Document has no extracted text. "
                "Re-upload the document or check if it is a scanned image that failed OCR."
            )

        # ------------------------------------------------------------------
        # 2. Sanitise document text (GAP-001)
        # ------------------------------------------------------------------
        safe_text = sanitise_document_text(doc.ocr_text)

        # ------------------------------------------------------------------
        # 3. Create draft record in GENERATING state
        # ------------------------------------------------------------------
        draft = Draft(
            firm_id=uuid.UUID(firm_id),
            draft_type=DraftType.CASE_SYNOPSIS,
            status=DraftStatus.GENERATING,
            title=f"Case Synopsis — {doc.file_name}",
            content="",  # filled after LLM returns
            generated_by_user_id=uuid.UUID(user_id) if user_id else None,
        )
        db.add(draft)
        await db.flush()  # get draft.id before LLM call

        logger.info(
            f"Generating synopsis draft={draft.id} document={document_id} firm={firm_id}"
        )

        # ------------------------------------------------------------------
        # 4. Call LLM with 90s timeout
        # ------------------------------------------------------------------
        try:
            raw_response = await asyncio.wait_for(
                self._llm.call_completion(
                    system_prompt=case_synopsis_prompt.system_prompt,
                    user_prompt=case_synopsis_prompt.format_user_prompt(
                        document_text=safe_text
                    ),
                    model=ModelType.GPT4O_MINI,
                    temperature=0.0,
                    max_tokens=case_synopsis_prompt.max_tokens,
                    firm_id=firm_id,
                    timeout=90,
                ),
                timeout=90.0
            )
        except asyncio.TimeoutError:
            logger.error(f"AI request timed out after 90s for synopsis document={document_id}")
            raise ValueError("AI request timed out — please try again")
        except Exception as exc:
            # Mark draft as failed so user can see it
            draft.status = DraftStatus.REJECTED
            draft.content = json.dumps({"error": str(exc)})
            await db.commit()
            logger.error(f"LLM failed for synopsis draft={draft.id}: {exc}")
            raise RuntimeError(f"Synopsis generation failed: {exc}") from exc

        # ------------------------------------------------------------------
        # 5. Parse and store result
        # ------------------------------------------------------------------
        try:
            parsed = _parse_synopsis_json(raw_response)
        except ValueError as exc:
            draft.status = DraftStatus.REJECTED
            draft.content = raw_response[:5000]  # store raw for debugging (no PII log)
            await db.commit()
            raise RuntimeError(str(exc)) from exc

        draft.status = DraftStatus.GENERATED
        draft.title = _build_synopsis_title(parsed)
        draft.content = json.dumps(parsed, ensure_ascii=False)
        draft.quality_score = int(parsed.get("confidence", 0)) * 10  # 0-10 → 0-100
        await db.commit()
        await db.refresh(draft)

        logger.info(f"Synopsis generated draft={draft.id} confidence={parsed.get('confidence')}")

        return SynopsisResult(
            synopsis_id=str(draft.id),
            document_id=document_id,
            case_name=parsed["case_name"],
            petitioner=parsed["petitioner"],
            respondent=parsed["respondent"],
            court=parsed["court"],
            judgment_date=parsed["judgment_date"],
            case_number=parsed["case_number"],
            facts=parsed["facts"],
            issues=parsed["issues"] or [],
            held=parsed["held"],
            citations_used=parsed["citations_used"] or [],
            relief_granted=parsed["relief_granted"],
            confidence=int(parsed.get("confidence", 0)),
            created_at=draft.created_at,
        )

    async def get_synopsis(
        self,
        db: AsyncSession,
        synopsis_id: str,
        firm_id: str,
    ) -> SynopsisResult:
        """
        Retrieve a previously generated synopsis by its draft ID.

        Raises:
            ValueError: Not found or wrong firm
        """
        stmt = select(Draft).where(
            Draft.id == uuid.UUID(synopsis_id),
            Draft.draft_type == DraftType.CASE_SYNOPSIS,
            Draft.deleted_at.is_(None),
        )
        result = await db.execute(stmt)
        draft: Optional[Draft] = result.scalar_one_or_none()

        if draft is None or str(draft.firm_id) != firm_id:
            raise ValueError("Synopsis not found")

        if draft.status == DraftStatus.GENERATING:
            raise ValueError("Synopsis is still being generated. Please try again shortly.")

        try:
            parsed = json.loads(draft.content)
        except (json.JSONDecodeError, TypeError):
            raise ValueError("Synopsis data is corrupted. Please regenerate.")

        return SynopsisResult(
            synopsis_id=str(draft.id),
            document_id="",  # not stored on draft — acceptable for retrieval
            case_name=parsed.get("case_name"),
            petitioner=parsed.get("petitioner"),
            respondent=parsed.get("respondent"),
            court=parsed.get("court"),
            judgment_date=parsed.get("judgment_date"),
            case_number=parsed.get("case_number"),
            facts=parsed.get("facts"),
            issues=parsed.get("issues") or [],
            held=parsed.get("held"),
            citations_used=parsed.get("citations_used") or [],
            relief_granted=parsed.get("relief_granted"),
            confidence=int(parsed.get("confidence", 0)),
            created_at=draft.created_at,
        )

    async def generate_synopsis_from_bytes(
        self,
        file_bytes: bytes,
        file_ext: str,
        firm_id: str,
    ) -> "SynopsisResult":
        """Generate synopsis directly from uploaded file bytes — no DB document required."""
        ft = file_ext.lower().lstrip(".")
        try:
            if ft == "pdf":
                from pypdf import PdfReader
                reader = PdfReader(BytesIO(file_bytes))
                raw_text = "\n".join(page.extract_text() or "" for page in reader.pages).strip()
                if not raw_text:
                    raise ValueError("PDF appears to be a scanned image — no selectable text found.")
            elif ft == "docx":
                from docx import Document as DocxDocument
                doc_obj = DocxDocument(BytesIO(file_bytes))
                raw_text = "\n".join(p.text for p in doc_obj.paragraphs).strip()
            else:
                raise ValueError(f"Unsupported file type '{ft}'. Supported: pdf, docx.")
        except ValueError:
            raise
        except Exception as exc:
            raise RuntimeError(f"Could not read file: {exc}") from exc

        # Truncate to ~14000 chars — legal notices can run 8-10 pages and
        # key demands/relief (Para 20+) and citations (Art 21, Rule 133A) appear
        # in later paragraphs. 8000 chars only covered ~first 4-5 pages.
        safe_text = sanitise_document_text(raw_text)[:14000]

        try:
            raw_response = await asyncio.wait_for(
                self._llm.call_completion(
                    system_prompt=case_synopsis_prompt.system_prompt,
                    user_prompt=case_synopsis_prompt.format_user_prompt(document_text=safe_text),
                    model=ModelType.GPT4O_MINI,
                    temperature=0.0,
                    max_tokens=1800,
                    firm_id=firm_id,
                    timeout=90,
                ),
                timeout=90.0,
            )
        except asyncio.TimeoutError:
            raise ValueError("AI request timed out — please try again")
        except Exception as exc:
            raise RuntimeError(f"Synopsis generation failed: {exc}") from exc

        parsed = _parse_synopsis_json(raw_response)

        # Validate document type — synopsis only works for court documents
        _validate_document_type(parsed)

        return SynopsisResult(
            synopsis_id=str(uuid.uuid4()),
            document_id="",
            case_name=parsed["case_name"],
            petitioner=parsed["petitioner"],
            respondent=parsed["respondent"],
            court=parsed["court"],
            judgment_date=parsed["judgment_date"],
            case_number=parsed["case_number"],
            facts=parsed["facts"],
            issues=parsed["issues"] or [],
            held=parsed["held"],
            citations_used=parsed["citations_used"] or [],
            relief_granted=parsed["relief_granted"],
            confidence=int(parsed.get("confidence", 0)),
            created_at=datetime.utcnow(),
        )

    async def export_synopsis_docx(
        self,
        db: AsyncSession,
        synopsis_id: str,
        firm_id: str,
    ) -> tuple[bytes, str]:
        """
        Export a synopsis as a formatted .docx file.

        Returns:
            (file_bytes, suggested_filename)

        Raises:
            ValueError: Not found, wrong firm, or not yet generated
        """
        from docx import Document as DocxDocument
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        synopsis = await self.get_synopsis(db, synopsis_id, firm_id)

        doc = DocxDocument()

        # --- Title ---
        title_para = doc.add_heading("CASE SYNOPSIS", level=1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph("")  # spacer

        # --- Helper to add labelled row ---
        def add_field(label: str, value: Optional[str]):
            if not value:
                return
            para = doc.add_paragraph()
            run_label = para.add_run(f"{label}: ")
            run_label.bold = True
            para.add_run(value)

        # --- Party/court info ---
        add_field("Case Name", synopsis.case_name)
        add_field("Case Number", synopsis.case_number)
        add_field("Petitioner", synopsis.petitioner)
        add_field("Respondent", synopsis.respondent)
        add_field("Court", synopsis.court)
        add_field("Date of Judgment", synopsis.judgment_date)

        doc.add_paragraph("")

        # --- Facts ---
        if synopsis.facts:
            doc.add_heading("Facts", level=2)
            doc.add_paragraph(synopsis.facts)

        # --- Issues ---
        if synopsis.issues:
            doc.add_heading("Issues", level=2)
            for i, issue in enumerate(synopsis.issues, 1):
                doc.add_paragraph(f"{i}. {issue}", style="List Number")

        # --- Held ---
        if synopsis.held:
            doc.add_heading("Held", level=2)
            doc.add_paragraph(synopsis.held)

        # --- Relief Granted ---
        if synopsis.relief_granted:
            doc.add_heading("Relief Granted", level=2)
            doc.add_paragraph(synopsis.relief_granted)

        # --- Citations Used ---
        if synopsis.citations_used:
            doc.add_heading("Citations Used", level=2)
            for citation in synopsis.citations_used:
                doc.add_paragraph(f"• {citation}")

        # --- Footer note ---
        doc.add_paragraph("")
        footer = doc.add_paragraph(
            f"Generated by Nikhar | Confidence Score: {synopsis.confidence}/10"
        )
        footer.runs[0].font.size = Pt(9)
        footer.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        # --- Serialise to bytes ---
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        file_bytes = buffer.read()

        safe_name = re.sub(r"[^\w\s-]", "", synopsis.case_name or "synopsis")
        safe_name = re.sub(r"\s+", "_", safe_name.strip())[:60]
        filename = f"synopsis_{safe_name}.docx"

        logger.info(f"Exported synopsis docx synopsis_id={synopsis_id} firm={firm_id}")

        return file_bytes, filename


# ---------------------------------------------------------------------------
# Singleton
# ---------------------------------------------------------------------------

_synopsis_service: Optional[SynopsisService] = None


def get_synopsis_service() -> SynopsisService:
    global _synopsis_service
    if _synopsis_service is None:
        _synopsis_service = SynopsisService(llm_service=get_llm_service())
    return _synopsis_service
